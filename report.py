import re
import requests
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from reportlab.platypus import SimpleDocTemplate, Spacer, Paragraph, Table, TableStyle, Image, PageBreak
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
import streamlit as st

def create_cover_page(doc, report_title):
    section = doc.sections[0]
    section.page_height = Pt(842)  # A4 size height
    section.page_width = Pt(595)   # A4 size width

    # Set margins to 2 cm (56.7 points)
    for section in doc.sections:
        section.top_margin = Pt(56.7)
        section.bottom_margin = Pt(56.7)
        section.left_margin = Pt(56.7)
        section.right_margin = Pt(56.7)

    # Add logo from URL to the cover page
    logo_url = "https://dcxsea.com/asset/images/logo/LOGO_DCX.png"
    response = requests.get(logo_url)
    if response.status_code == 200:
        logo_image = BytesIO(response.content)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(logo_image, width=Pt(150))  # Adjust width as needed

        # Add company name after the logo
        run = paragraph.add_run("  DCx Co., Ltd.")
        run.font.size = Pt(24)
        run.bold = True
    else:
        st.error("Failed to download the logo image.")

    doc.add_paragraph("\n")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = cover.add_run("Report")
    run.font.size = Pt(20)
    run.bold = True
    cover.add_run("\n").font.size = Pt(24)

    # Add some spacing
    doc.add_paragraph("\n")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = cover.add_run("Indigenous Agriculture Adaptation")
    run.font.size = Pt(24)
    run.bold = True
    cover.add_run("\n").font.size = Pt(24)

    for _ in range(2):
        doc.add_paragraph("\n")

    # Add prepared for
    prepared_for = doc.add_paragraph()
    prepared_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prepared_for.add_run("Prepared for: Jack Jasmin")
    run.font.size = Pt(14)
    run.add_break()  # Add a break line
    run.add_break()
    run = prepared_for.add_run("Prepared by: Black Eye Team")
    run.font.size = Pt(14)
    run.add_break()
    run.add_break()

    # Place the date at the bottom center of the cover page
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(14)

    # Create a new section for the rest of the document, so the footer with the date is not repeated
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Ensure the new section starts with a fresh header and footer
    new_section.footer.is_linked_to_previous = False
    new_section.footer.paragraphs[0].clear()
    new_section.header.is_linked_to_previous = False

def add_markdown_formatted_text(paragraph, text):
    bold_italic_pattern = re.compile(r'\*\*\*(.+?)\*\*\*')
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')
    italic_pattern = re.compile(r'\*(.+?)\*')
    patterns = [
        ('bold_italic', bold_italic_pattern),
        ('bold', bold_pattern),
        ('italic', italic_pattern)
    ]

    def replace_match(match, style):
        if style == 'bold_italic':
            run = paragraph.add_run(match.group(1))
            run.bold = True
            run.italic = True
        elif style == 'bold':
            run = paragraph.add_run(match.group(1))
            run.bold = True
        elif style == 'italic':
            run = paragraph.add_run(match.group(1))
            run.italic = True

    cursor = 0
    text_length = len(text)
    while cursor < text_length:
        nearest_match = None
        nearest_style = None
        nearest_start = text_length

        for style, pattern in patterns:
            match = pattern.search(text, cursor)
            if match:
                start = match.start()
                if start < nearest_start:
                    nearest_match = match
                    nearest_style = style
                    nearest_start = start

        if nearest_match:
            if nearest_start > cursor:
                paragraph.add_run(text[cursor:nearest_start])
            replace_match(nearest_match, nearest_style)
            cursor = nearest_match.end()
        else:
            paragraph.add_run(text[cursor:])
            break

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def save_report_as_word(report, filename):
    try:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Pt(56.7)
            section.bottom_margin = Pt(56.7)
            section.left_margin = Pt(56.7)
            section.right_margin = Pt(56.7)

        create_cover_page(doc, filename.split('.')[0])

        lines = report.split('\n')
        table = None
        lines = lines[1:]
        # table = None

        for line in lines:
            # line = line.strip()
            # if not line:
            #     continue 


            if line.strip() == "---":
                doc.add_page_break()
            elif line.strip().startswith("# "):
                paragraph = doc.add_heading(line.strip()[2:], level=1)
            elif line.strip().startswith("## "):
                paragraph = doc.add_heading(line.strip()[3:], level=2)
            elif line.strip().startswith("### "):
                paragraph = doc.add_heading(line.strip()[4:], level=3)
            elif line.strip().startswith("#### "):
                paragraph = doc.add_heading(line.strip()[5:], level=4)
            elif line.strip().startswith("##### "):
                paragraph = doc.add_heading(line.strip()[6:], level=5)
            elif line.strip().startswith("###### "):
                paragraph = doc.add_heading(line.strip()[7:], level=6)
            elif line.strip().startswith("* "):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif "|" in line:
                table_data = [cell.strip() for cell in line.split('|') if cell]
                
                if not table:
                    table = doc.add_table(rows=1, cols=len(table_data))
                    hdr_cells = table.rows[0].cells
                    for i, cell_data in enumerate(table_data):
                        hdr_cells[i].text = cell_data
                else:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(table_data):
                        row_cells[i].text = cell_data
            else:
                paragraph = doc.add_paragraph()
                add_markdown_formatted_text(paragraph, line)
                table = None

        doc.save(filename)
    except Exception as e:
        st.error(f"Failed to save Word report: {e}")


def save_report_as_pdf(report, filename):
    try:
        doc = SimpleDocTemplate(filename, pagesize=letter)
        elements = []

        # Adjust the logo path and size
        logo_path = "https://dcxsea.com/asset/images/logo/LOGO_DCX.png"
        logo = Image(logo_path, 2 * inch, 1 * inch)

        # Create the company name Paragraph style and the paragraph itself
        company_name_style = ParagraphStyle(name='CompanyNameStyle',fontName='Helvetica-Bold', fontSize=20, leading=22, alignment=TA_CENTER)
        company_name = Paragraph("DCx Co., Ltd.", company_name_style)

        # Create a table to hold the logo and the company name side by side
        logo_and_name_table = Table(
            [[logo, company_name]],
            colWidths=[2 * inch, 2 * inch],  # Adjust column widths as needed
        )

        # Style the table to center everything
        logo_and_name_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center both the logo and the company name
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Vertically align them in the middle
        ]))

        # Center the table itself on the page
        elements.append(Spacer(1, 10))  # Add space at the top to center the logo and name vertically
        elements.append(logo_and_name_table)
        elements.append(Spacer(1, 36))

        # Add the cover page details
        cover_style = ParagraphStyle(name='CoverStyle', fontSize=24, alignment=TA_CENTER, spaceAfter=12)
        elements.append(Paragraph("Report", cover_style))
        elements.append(Spacer(1, 40))
        elements.append(Paragraph("Indigenous Agriculture Adaptation", cover_style))
        elements.append(Spacer(1, 60))
        elements.append(Paragraph(f"Prepared for: Jack Jasmin", ParagraphStyle(name='Normal', fontSize=14, alignment=TA_CENTER)))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Prepared by: Black Eye Team", ParagraphStyle(name='Normal', fontSize=14, alignment=TA_CENTER)))
        elements.append(Spacer(1, 300))
        elements.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", ParagraphStyle(name='Normal', fontSize=14, alignment=TA_CENTER)))
        elements.append(Spacer(1, 10))

        # Create the styles for the document
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(name='Normal', fontSize=12, leading=14, alignment=TA_JUSTIFY)
        heading1_style = ParagraphStyle(name='Heading1', fontSize=16, leading=18, alignment=TA_LEFT, spaceAfter=12)
        heading2_style = ParagraphStyle(name='Heading2', fontSize=14, leading=16, alignment=TA_LEFT, spaceAfter=12)
        heading3_style = ParagraphStyle(name='Heading3', fontSize=12, leading=14, alignment=TA_LEFT, spaceAfter=12)
        heading4_style = ParagraphStyle(name='Heading4', fontSize=10, leading=12, alignment=TA_LEFT, spaceAfter=12)
        bullet_style = ParagraphStyle(name='Bullet', fontSize=12, leading=14, alignment=TA_JUSTIFY, bulletIndent=12)

        # Split the report into lines
        lines = report.split('\n')
        lines = lines[1:]  # Skip the first line

        # Convert the report text into Paragraphs for the PDF
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 12))
            elif line == "---":
                elements.append(PageBreak()) 
            elif line.startswith('# '):
                elements.append(Paragraph(line[2:], heading1_style))
            elif line.startswith('## '):
                elements.append(Paragraph(line[3:], heading2_style))
            elif line.startswith('### '):
                elements.append(Paragraph(line[4:], heading3_style))
            elif line.startswith('#### '):
                elements.append(Paragraph(line[5:], heading4_style))
            elif line.startswith('**** '):
                elements.append(Paragraph(line[5:], bullet_style))
            elif line.startswith('*** '):
                elements.append(Paragraph(line[4:], bullet_style))
            elif line.startswith('** '):
                elements.append(Paragraph(line[3:], bullet_style))
            elif line.startswith('* '):
                elements.append(Paragraph(line[2:], bullet_style))
            elif "|" in line:
                # Handle table creation more robustly
                table_data = [[Paragraph(cell.strip(), normal_style) for cell in line.split('|') if cell]]
                
                table = Table(table_data, colWidths=[1.5 * inch] * len(table_data[0]))  # Adjust colWidths based on content
                
                table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTSIZE', (0, 0), (-1, -1), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                ]))
                elements.append(table)
            else:
                # Handle bold text with '**' in PDF
                bold_pattern = re.compile(r'\*\*(.+?)\*\*')
                if bold_pattern.search(line):
                    parts = bold_pattern.split(line)
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            elements.append(Paragraph(part, ParagraphStyle(name='Bold', fontSize=12, leading=14, alignment=TA_JUSTIFY, fontName='Helvetica-Bold')))
                        else:
                            elements.append(Paragraph(part, normal_style))
                else:
                    elements.append(Paragraph(line, normal_style))
            elements.append(Spacer(1, 12))

        # Build the PDF
        doc.build(elements)
    except Exception as e:
        st.error(f"Failed to save PDF report: {e}")


